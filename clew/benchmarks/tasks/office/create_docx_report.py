"""Office task — create a .docx with a heading and 2 paragraphs.

Category: office
Section: office
Difficulty: easy

The agent must use ``office_create`` + ``office_add_heading`` +
``office_add_paragraph`` (or write_file with python-docx) to produce
a ``report.docx`` containing:
- A heading "Quarterly Summary"
- A paragraph starting with "Revenue grew"
- A paragraph starting with "Costs decreased"

Pass criteria: ``report.docx`` exists and is a valid .docx; opening
it with python-docx yields the expected heading + 2 paragraphs.
"""

from __future__ import annotations

from pathlib import Path

from ..._base import (
    Difficulty,
    EvaluationReport,
    Section,
    TaskSpec,
)


def setup(workspace: str) -> None:
    root = Path(workspace)
    root.mkdir(parents=True, exist_ok=True)
    (root / "README.md").write_text(
        "# Office workspace\n\nCreate report.docx here.\n",
        encoding="utf-8",
    )


def evaluate(workspace: str, agent_output: str, tool_calls: list) -> EvaluationReport:
    root = Path(workspace)
    docx = root / "report.docx"
    criteria = []

    # 1. report.docx exists.
    c1 = docx.is_file()
    criteria.append({"name": "report.docx exists", "passed": c1})

    # 2. It's a valid .docx — open it with python-docx.
    c2 = False
    paras: list = []
    if c1:
        try:
            from docx import Document
            doc = Document(str(docx))
            paras = list(doc.paragraphs)
            c2 = True
        except Exception:
            c2 = False
    criteria.append({"name": "report.docx is a valid .docx", "passed": c2})

    # 3. Has a heading containing "Quarterly Summary".
    c3 = False
    if c2:
        for p in paras:
            if p.style and "Heading" in (p.style.name or ""):
                if "quarterly summary" in (p.text or "").lower():
                    c3 = True
                    break
    criteria.append({
        "name": "heading 'Quarterly Summary' present",
        "passed": c3,
    })

    # 4. Has paragraphs starting with "Revenue grew" and "Costs decreased".
    text_blob = "\n".join(p.text for p in paras) if c2 else ""
    c4a = "revenue grew" in text_blob.lower()
    c4b = "costs decreased" in text_blob.lower()
    c4 = c4a and c4b
    criteria.append({
        "name": "contains 'Revenue grew' and 'Costs decreased' paragraphs",
        "passed": c4,
    })

    # 5. (Informational) Agent used office_* tools.
    office_calls = [
        tc for tc in tool_calls
        if str(tc.get("name", "")).startswith("office_")
    ]
    c5 = len(office_calls) >= 1
    criteria.append({
        "name": "agent used office_* tools",
        "passed": c5,
    })

    mandatory = [c["passed"] for c in criteria[:4]]
    passed = all(mandatory)
    return EvaluationReport(
        passed=passed,
        reason="all mandatory criteria met" if passed else "mandatory criteria failed",
        details=f"office tool calls: {len(office_calls)}",
        checked_criteria=criteria,
    )


def build() -> TaskSpec:
    return TaskSpec(
        id="office_create_docx_report",
        section=Section.OFFICE,
        difficulty=Difficulty.EASY,
        description="Create report.docx with heading + 2 paragraphs using office tools.",
        prompt=(
            "Create a new Word document called report.docx in the "
            "workspace root. It should contain: (1) a heading 'Quarterly "
            "Summary', (2) a paragraph starting with 'Revenue grew by 12% "
            "year over year', (3) a paragraph starting with 'Costs "
            "decreased by 5% due to vendor renegotiation'. Use the office_* "
            "tools (office_create, office_add_heading, office_add_paragraph)."
        ),
        setup=setup,
        evaluate=evaluate,
        expected_duration_s=30.0,
        tags=["office", "docx"],
        min_iterations=6,
        max_time_s=240.0,
    )
